import streamlit as st
import tempfile
import os
from pathlib import Path
import docx 
import pymupdf
from IA.Anthropic_connection import *
from IA.OpenIA_connection import *
from surligner import *
from conversion import *
from prompt import *

# ─────────────────────────────────────────────
# Interface Streamlit
# ─────────────────────────────────────────────

st.set_page_config(
    page_title="Détection d'annales",
    page_icon="📄",
    layout="centered"
)

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if st.session_state.authenticated :

    st.title("📄 Détection d'annales dans un ronéo")
    st.caption("Importe tes fichiers, réponds aux questions, récupère le résultat.")

    # ── Étape 1 : Upload des fichiers ──────────────────────────────────────────

    st.header("1 · Fichiers sources")

    col1, col2 = st.columns(2)

    with col1:
        roneo_file = st.file_uploader(
            "Ronéo source",
            type=["docx"],
            help="Le ronéo dans lequel détecter les annales"
        )

    with col2:
        annales_file = st.file_uploader(
            "Fichier annales",
            type=["pdf"],
            help="Le fichier annales à enrichir / compléter"
        )

    # ── Étape 2 : Analyse + questions dynamiques ──────────────────────────────

    # Initialiser la variable de session pour le bouton
    if "bt_demarrage" not in st.session_state:
        st.session_state.bt_demarrage = False

    # Afficher le bouton seulement si les deux fichiers sont chargés
    if roneo_file and annales_file:
        
        # SAUVEGARDER LES NOMS DES FICHIERS
        if "roneo_file_name" not in st.session_state:
            st.session_state.roneo_file_name = roneo_file.name
        if "annales_file_name" not in st.session_state:
            st.session_state.annales_file_name = annales_file.name
        
        # SAUVEGARDER LE CONTENU DES FICHIERS
        if "roneo_file_bytes" not in st.session_state:
            st.session_state.roneo_file_bytes = roneo_file.getvalue()
        if "annales_file_bytes" not in st.session_state:
            st.session_state.annales_file_bytes = annales_file.getvalue()
        
        st.session_state.bt_demarrage = True
    else:
        st.info("Importe les deux fichiers pour démarrer l'analyse.", icon="ℹ️")


    if st.session_state.bt_demarrage:
        st.caption("Le Script est en cours de fonctionnement, cela peut prendre 2/3min")
        #st.header("2 · Questions")

         # ✅ VÉRIFIER QUE LES DONNÉES SONT EN SESSION
        if "roneo_file_bytes" not in st.session_state or "annales_file_bytes" not in st.session_state or "roneo_file_name" not in st.session_state or "annales_file_name" not in st.session_state:
            st.error("⚠️ Fichiers manquants. Recharge la page et réimporte-les.")
            st.session_state.bt_demarrage = False
            st.stop()

        # Sauvegarde temporaire des fichiers uploadés
        with tempfile.TemporaryDirectory() as tmpdir:
            # ✅ UTILISER LES DONNÉES SAUVEGARDÉES
            roneo_path = os.path.join(tmpdir, st.session_state.roneo_file_name)
            annales_path = os.path.join(tmpdir, st.session_state.annales_file_name)
            
            with open(roneo_path, "wb") as f:
                f.write(st.session_state.roneo_file_bytes)
            with open(annales_path, "wb") as f:
                f.write(st.session_state.annales_file_bytes)

            # Lance l'analyse (mise en cache pour éviter de ré-analyser à chaque interaction)
            @st.cache_data(show_spinner="Analyse en cours…")
            
            # --- Transformer le cours docx en cours txt ---
            def remove_espace(txt):
                #supprime les trucs en trop
                t = txt.split('\n')
                r=[]
                for i in t:
                    if i!="":
                        r.append(i)
                return "\n".join(r)
            
            def docx_to_txt(source_path,result_name):
                fichier_cours = docx.Document(source_path)
                cours_txt = ""
                for p in fichier_cours.paragraphs:
                    cours_txt+="\n" 
                    cours_txt += p.text
                for t in fichier_cours.tables:
                    for r in t.rows:
                        for cell in r.cells:
                            cours_txt+="\n" 
                            cours_txt += cell.text
                cours_txt = remove_espace(cours_txt)
                #print(cours_txt)
                result_path = os.path.join(tmpdir, result_name)
                with open(result_path, "w", encoding="utf-8") as f2:
                    f2.write(cours_txt)
                return result_path

            roneo_txt_path = docx_to_txt(roneo_path,"roneo_txt.txt")
            print(roneo_txt_path)

            # --- Transformer les annales pdf en annales txt ---
            def pdf_to_txt(source_path,result_name):
                doc = pymupdf.open(source_path) # open a document
                txt_annales = ""
                for page in doc: # iterate the document pages
                    text = page.get_text() #recupere le text format utf8
                    txt_annales += text
                #print(txt_annales)
                result_path = os.path.join(tmpdir, result_name)
                with open(result_path, "w", encoding="utf-8") as f2:
                    f2.write(txt_annales)
                return result_path
            
            annales_txt_path = pdf_to_txt(annales_path,"oki")
            print(annales_txt_path)

            
            # Connection à l'IA
            apikey_openia = st.secrets["API_KEY_OPENAI"]
            apikey_anthropic = st.secrets["API_KEY_ANTHROPIC"]

            st.title("⚙️ Configuration IA")

            # 1. Choix du provider
            provider = st.selectbox(
                "Choisis le type d'IA :",
                ["ChatGPT", "Anthropic (Claude)"]
            )

            # 2. Choix du modèle selon provider
            if provider == "ChatGPT":
                model = st.selectbox(
                    "Choisis le modèle ChatGPT :",
                    [
                        "gpt-5.4",
                        "gpt-5.4-mini",
                        "gpt-5.4-nano",
                        "gpt-5.2",
                        "gpt-5.2-pro",
                        "gpt-5.1",
                        "gpt-5",
                        "gpt-5-mini",
                        "gpt-5-nano",
                        "gpt-4.1",
                        "gpt-4.1-mini",
                        "gpt-4.1-nano"
                    ]
                )
            else:
                model = st.selectbox(
                    "Choisis le modèle Claude :",
                    [
                        "claude-opus-4-6",
                        "claude-opus-4-5",
                        "claude-opus-4-1",
                        "claude-opus-4",
                        "claude-sonnet-4-6",
                        "claude-sonnet-4",
                        "claude-sonnet-3-7",
                        "claude-haiku-4-5",
                        "claude-haiku-3-5",
                        "claude-haiku-3"
                    ]
                )

            # 3. Zone de prompt (modifiable)
            pt = """Tu as pour mission d'identifier dans les cours (fichier cours.txt) les connaissances requises pour répondre annales (annales.txt).

            Fichiers :
            - Un fichier annales.txt contenant les annales tombées au concours et leur correction
            - Un fichier cours.txt contenant un cours dans lequel on cherche les notions demandé le jour de l'examen

            Consignes :
            - Rescencer dans le cours des mots ou bout de phrase essentiel à connaître pour répondre aux annales.
            - Relever les passages exactes du cours

            Format de sortie :
            Liste courte, chaque élément = 1 notion (3 à 6 mots max qui correspondent exactement à ce qui est écrit dans cours.txt) avec, pour faciliter le repérage, le mot précédant + [AN] + notion + [/AN] + mot suivant. J'attends une liste sans ajout autre.
            """
            prompt = st.text_area(
                "Prompt :",
                value=PROMPT_UNIQUE_AMELIORE,
                height=300
            )

            # 4. Bouton valider
            if st.button("🚀 Valider"):
                st.write(f"[{model}]")
                if provider == "ChatGPT":
                    #Upload file
                    annales_txt_id = IA_upload_openIA(apikey_openia, annales_txt_path)
                    roneo_txt_id = IA_upload_openIA(apikey_openia, roneo_txt_path)

                    #Ask IA
                    reponse_ia = IA_ask_openIA(apikey_openia,prompt,[annales_txt_id,roneo_txt_id],model)
                else :
                    #Upload file
                    annales_txt_id = IA_upload_anthropic(apikey_anthropic, "annales.txt",annales_txt_path)
                    roneo_txt_id = IA_upload_anthropic(apikey_anthropic, "roneo.txt",roneo_txt_path)

                    #Ask IA
                    reponse_ia = IA_ask_anthropic(apikey_anthropic,prompt,[annales_txt_id,roneo_txt_id],model)

                st.success("Configuration validée !")

                st.write(reponse_ia)

                #Traitement de la réponse de l'IA :
                def creation_notion_ls(reponse_ia):
                    notions_ls = []
                    for i in reponse_ia.split("\n"):
                        j = i.split('[AN]')
                        if len(j)>1:
                            notion = j[1].split('[/AN]')[0]
                            #retire les espaces si jamais au début et à la fin
                            if notion[0] == " ": notion = notion[1:]
                            if notion[-1] == " ": notion = notion[:-1]
                            notions_ls.append(notion)
                    return notions_ls

                print(reponse_ia)
                
                notion_ls = creation_notion_ls(reponse_ia)
                st.write(notion_ls)


                # Stocker les données
                st.session_state.notion_ls = notion_ls
                st.session_state.index = 0
                st.session_state.selection = []
                st.session_state.ia_done = True  # Flag pour indiquer que l'IA a terminé
            
            # 🔹 SECTION DE TRI (en dehors du bouton Valider)
            if "ia_done" in st.session_state and st.session_state.ia_done:
                st.title("🧠 Tri des notions")

                # 🔹 Fin du processus
                if st.session_state.index >= len(st.session_state.notion_ls):
                    st.success("✅ Tri terminé !")

                    st.write("### 📌 Liste conservée :")
                    st.write(st.session_state.selection)

                    #Surligner dans le RONEO
                    roneo_final_path = os.path.join(tmpdir, f"NEW_{roneo_file.name}")
                    surligner_mots(roneo_path, st.session_state.selection, roneo_final_path)

                    # Lire le fichier en binaire
                    with open(roneo_final_path, "rb") as f:
                        data = f.read()

                    # Bouton de téléchargement
                    st.download_button(
                        label="💾 Télécharger le fichier Word",
                        data=data,
                        file_name=f"NEW_{roneo_file.name}",  # Nom que verra l'utilisateur
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


                else:
                    terme = st.session_state.notion_ls[st.session_state.index]

                    st.write(f"### Terme : {terme}")

                    col1, col2 = st.columns(2)

                    # 🔴 Bouton supprimer
                    with col1:
                        if st.button("❌ Je supprime"):
                            st.session_state.index += 1
                            st.rerun()

                    # 🟢 Bouton garder
                    with col2:
                        if st.button("✅ Je garde"):
                            st.session_state.selection.append(terme)
                            st.session_state.index += 1
                            st.rerun()



else :
    #Mot de passe
    st.title("🔐 Accès protégé")
    
    code = st.text_input("Entre le code secret :")

    if st.button("Valider"):
        if code == st.secrets['SECRET_CODE']:
            st.session_state.authenticated = True
            st.rerun()  # recharge l'app
        else:
            st.error("Code incorrect ❌")


