from docx import Document
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

"""
Structure d'un paragraphe
<w:p>
  <w:r><w:t>Voici une équation : </w:t></w:r>
  <m:oMath> ... </m:oMath>
  <w:r><w:t> et du texte après</w:t></w:r>
</w:p>
"""


def __analyser_paragraph(paragraph, mot_ls):
    """
    <CT_R '<w:r>' at 0x2381a0a0470> {http://schemas.openxmlformats.org/wordprocessingml/2006/main}r
    <Element {http://schemas.openxmlformats.org/officeDocument/2006/math}oMath at 0x2381a097700> {http://schemas.openxmlformats.org/officeDocument/2006/math}oMath  
    """
    text_total = ""
    for elem in paragraph._element:
        tag = elem.tag
        #print(f"{elem} {elem.tag}")
        if tag.endswith('}r'): #run
            
            texts = elem.findall('.//w:t', namespaces=elem.nsmap)
            image = elem.findall('.//w:drawing', namespaces=elem.nsmap)
            print(texts, image)
            if len(texts) == 0:
                #text_total += "[ELEMENT]"
                print("IMAGE")
            else:
                text_total = text_total + texts[0].text
        elif tag.endswith('}oMath') or tag.endswith('}oMathPara') :
            print("EQUATION")
            #text_total += "[ELEMENT]"

    print(f"Text total = {text_total}")

    #Détecter les positions des mots à surligner :
    position_cible = [] # [ [début, fin] ]
    for mot in mot_ls:
        spliter = text_total.lower().split(mot.lower())
        if len(spliter) > 1:
            spliter.pop()
            counter = 0
            for d in spliter:
                position_cible.append([counter+len(d),counter+len(d)+len(mot)])
                counter = counter + len(d) + len(mot)
    print(position_cible)

    #Surlignagne
    
    
    for cible in position_cible:
        counter = 0 #compteur des lettres 
        index_balise = 0 #index des balises XML
        for elem in paragraph._element:
            print(index_balise, len(list(paragraph._element)))
            tag = elem.tag
            #print(f"{elem} {elem.tag}")
            if tag.endswith('}r'): #run
                texts = elem.findall('.//w:t', namespaces=elem.nsmap)
                image = elem.findall('.//w:drawing', namespaces=elem.nsmap)
                print(texts, image)
                
                if len(texts) == 0 or len(image) != 0:
                    print("IMAGE")
                    #ON NE TOUCHE A RIEN
                else:
                    print(texts[0].text)
                    txt = texts[0].text
                
                    l=0
                    before=""
                    inside=""
                    after=""
                    for lettre in txt:
                        if counter + l < cible[0]:
                            before += lettre
                        elif counter + l >= cible[1]:
                            after += lettre
                        else:
                            inside += lettre
                        l+=1
                    print(f"BEFORE : {before} INSIDE : {inside} AFTER : {after}")
                    
                    #création des balises
                    #on crée autant de balise que de texte dispo (on les insert au bonne endroit)
                    nb_balises = sum(1 for i in [before, after, inside] if i) - 1 
                    print("COMPTAGE :" ,nb_balises)
                    balise_source = paragraph._element[index_balise]
                    
                    for i in range(nb_balises):
                        print('NOUVELLE BALISE')
                        new_elem = deepcopy(balise_source)
                        paragraph._element.insert(index_balise, new_elem)

                        """
                        print(paragraph._element[index_balise].xml)
                        print('-'*44)
                        print(paragraph._element[index_balise+1].xml)
                        exit(0)
                        """
                    
                    
                    if before != "":
                        #On garde le même index globale et on change le txt du before, car il est petu être tronqué pour inside
                        paragraph._element[index_balise].text = before
                        if nb_balises > 0:
                            index_balise += 1
                            nb_balises -= 1
                    
                    if inside != "":
                        # Surligne en jaune
                        elem = paragraph._element[index_balise]

                        # 1. récupérer ou créer <w:rPr>
                        rPr = elem.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            elem.insert(0, rPr)
                        
                        # 2. créer <w:highlight>
                        highlight = OxmlElement('w:highlight')
                        highlight.set(qn('w:val'), 'yellow')

                        # 3. ajouter dans les propriétés
                        rPr.append(highlight)

                        paragraph._element[index_balise].text = inside
                        if nb_balises > 0:
                            index_balise += 1
                            nb_balises -= 1
                    
                    if after != "" :
                        paragraph._element[index_balise].text = after
                        if nb_balises > 0:
                            index_balise += 1
                            nb_balises -= 1
                    
                    counter+=len(txt)
            elif tag.endswith('}oMath') or tag.endswith('}oMathPara') :
                print("EQUATION")
        
            
            index_balise += 1
        
        




def surligner_mots(doc, mot_ls, save_path):
    fichier = Document(doc)
    for paragraph in fichier.paragraphs:
        __analyser_paragraph(paragraph, mot_ls)
        
    
    for table in fichier.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    __analyser_paragraph(paragraph, mot_ls)

    fichier.save(save_path)


#surligner_mots("./v3/test.docx",["surligner", "un test"])
#surligner_mots("./ressources/UE6-C8-L1-Oussalah.docx",["2/3 du volume", "lobe gauche", "division chirurgicale", "bilirubine"],"./v5/result.docx")
#surligner_mots("./v5/test.docx", ["à tous", "vous"], "./v5/surlignage.docx")
