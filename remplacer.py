from docx import Document
from copy import deepcopy


def __analyser_paragraph_remplacement(paragraph, phrase_avant, phrase_apres):
    """
    Analyse un paragraphe et remplace phrase_avant par phrase_apres
    en préservant le formatage des runs.
    """
    txt = paragraph.text
    
    # Vérifier si la phrase à remplacer existe
    if phrase_avant not in txt:
        return False
    
    # Trouver la position
    debut_remplacement = txt.find(phrase_avant)
    if debut_remplacement == -1:
        return False
    
    fin_remplacement = debut_remplacement + len(phrase_avant)
    
    counter = 0
    result_run = []
    remplacement_fait = False
    
    for run in paragraph.runs:
        txt_run = run.text
        run_debut = counter
        run_fin = counter + len(txt_run)
        
        # Run avant la zone de remplacement
        if run_fin <= debut_remplacement:
            result_run.append(deepcopy(run))
            counter += len(txt_run)
            continue
        
        # Run après la zone de remplacement
        if run_debut >= fin_remplacement:
            result_run.append(deepcopy(run))
            counter += len(txt_run)
            continue
        
        # Run qui chevauche la zone de remplacement
        before = ""
        inside = ""
        after = ""
        
        for i, lettre in enumerate(txt_run):
            pos_absolue = counter + i
            
            if pos_absolue < debut_remplacement:
                before += lettre
            elif pos_absolue >= fin_remplacement:
                after += lettre
            else:
                inside += lettre
        
        # Créer les nouveaux runs
        if before:
            new_run = deepcopy(run)
            new_run.text = before
            result_run.append(new_run)
        
        # Remplacement (une seule fois)
        if inside and not remplacement_fait:
            new_run = deepcopy(run)
            new_run.text = phrase_apres
            result_run.append(new_run)
            remplacement_fait = True
        
        if after:
            new_run = deepcopy(run)
            new_run.text = after
            result_run.append(new_run)
        
        # Élément vide (image, etc.)
        if not before and not inside and not after:
            result_run.append(deepcopy(run))
        
        counter += len(txt_run)
    
    # Remplacer les runs
    for run in paragraph.runs:
        paragraph._element.remove(run._element)
    
    for new_run in result_run:
        paragraph._element.append(new_run._element)
    
    return True


def remplacer_phrases(doc_path, liste_remplacements, save_path):
    """
    Remplace une liste de phrases dans un document Word.
    
    Args:
        doc_path: Chemin du document source
        liste_remplacements: Liste [[phrase_avant, phrase_apres], ...]
        save_path: Chemin de sauvegarde
    
    Returns:
        dict: Statistiques
    """
    fichier = Document(doc_path)
    stats = {
        'total_remplacements': len(liste_remplacements),
        'remplacements_effectues': 0,
        'paragraphes_modifies': 0,
        'tableaux_modifies': 0
    }
    
    # Paragraphes normaux
    for paragraph in fichier.paragraphs:
        para_modifie = False
        for phrase_avant, phrase_apres in liste_remplacements:
            if __analyser_paragraph_remplacement(paragraph, phrase_avant, phrase_apres):
                stats['remplacements_effectues'] += 1
                para_modifie = True
        
        if para_modifie:
            stats['paragraphes_modifies'] += 1
    
    # Tableaux
    for table in fichier.tables:
        table_modifie = False
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for phrase_avant, phrase_apres in liste_remplacements:
                        if __analyser_paragraph_remplacement(paragraph, phrase_avant, phrase_apres):
                            stats['remplacements_effectues'] += 1
                            table_modifie = True
        
        if table_modifie:
            stats['tableaux_modifies'] += 1
    
    fichier.save(save_path)
    return stats